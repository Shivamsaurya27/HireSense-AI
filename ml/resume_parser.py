"""
ml/resume_parser.py
--------------------
Resume Parsing Module for HireSense AI.

Responsible for extracting clean, raw text from resume files supplied
in PDF, DOCX, or TXT format. This module performs NO NLP or scoring —
its only job is reliable, fault-tolerant text extraction that later
modules (text_processor, skill_extractor, ats_scorer, etc.) can build on.

Design goals:
- Never crash the pipeline on a bad/corrupted/empty file.
- Always return a predictable result object with a success flag,
  extracted text, and diagnostic info (errors/warnings, file metadata).
- Support multiple extraction backends for PDFs (pdfplumber primary,
  PyPDF2 fallback) since real-world resumes are inconsistently formatted.
"""

from __future__ import annotations

import os
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

# --- Third-party parsing libraries -----------------------------------------
# Each import is wrapped so that a missing optional dependency doesn't crash
# the whole module — it just disables that specific file-type support.
try:
    import pdfplumber
    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    _PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    _PYPDF2_AVAILABLE = True
except ImportError:
    _PYPDF2_AVAILABLE = False

try:
    import docx  # python-docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


logger = logging.getLogger("hiresense.resume_parser")
if not logger.handlers:
    # Basic default logging config so this module is usable standalone.
    logging.basicConfig(level=logging.INFO)


# Supported file extensions mapped to a human-readable label.
SUPPORTED_EXTENSIONS = {
    ".pdf": "PDF",
    ".docx": "DOCX",
    ".txt": "TXT",
}

# A resume with less than this many characters of extracted text is
# treated as effectively empty/unreadable (e.g. scanned image PDF with
# no OCR text layer).
MIN_VALID_TEXT_LENGTH = 30


@dataclass
class ParsedResume:
    """Structured result of a resume parsing attempt.

    This is the contract the rest of the pipeline depends on — always
    check `success` before trusting `raw_text`.
    """
    file_path: str
    file_type: Optional[str] = None
    raw_text: str = ""
    char_count: int = 0
    success: bool = False
    errors: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        """Convert to a plain dict (useful for JSON responses / logging)."""
        return {
            "file_path": self.file_path,
            "file_type": self.file_type,
            "raw_text": self.raw_text,
            "char_count": self.char_count,
            "success": self.success,
            "errors": self.errors,
            "warnings": self.warnings,
        }


class ResumeParser:
    """Extracts raw text from resume files (PDF, DOCX, TXT).

    Usage:
        parser = ResumeParser()
        result = parser.parse("resumes/john_doe.pdf")
        if result.success:
            print(result.raw_text)
        else:
            print(result.errors)
    """

    def __init__(self) -> None:
        # Surface missing optional dependencies once, at construction time,
        # rather than repeatedly during parsing.
        if not _PDFPLUMBER_AVAILABLE and not _PYPDF2_AVAILABLE:
            logger.warning(
                "Neither pdfplumber nor PyPDF2 is installed. "
                "PDF parsing will not be available."
            )
        if not _DOCX_AVAILABLE:
            logger.warning(
                "python-docx is not installed. DOCX parsing will not be available."
            )

    # -- Public API -----------------------------------------------------

    def parse(self, file_path: str) -> ParsedResume:
        """Parse a resume file and return a ParsedResume result.

        This method never raises — all failure modes are captured inside
        the returned ParsedResume (success=False, errors populated).
        """
        result = ParsedResume(file_path=file_path)

        # --- Basic file validation --------------------------------------
        path = Path(file_path)

        if not path.exists():
            result.errors.append(f"File not found: {file_path}")
            return result

        if not path.is_file():
            result.errors.append(f"Path is not a file: {file_path}")
            return result

        if path.stat().st_size == 0:
            result.errors.append("File is empty (0 bytes).")
            return result

        extension = path.suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            result.errors.append(
                f"Unsupported file type '{extension}'. "
                f"Supported types: {', '.join(SUPPORTED_EXTENSIONS.values())}"
            )
            return result

        result.file_type = SUPPORTED_EXTENSIONS[extension]

        # --- Dispatch to the correct extractor --------------------------
        try:
            if extension == ".pdf":
                text, warnings = self._extract_pdf_text(file_path)
            elif extension == ".docx":
                text, warnings = self._extract_docx_text(file_path)
            elif extension == ".txt":
                text, warnings = self._extract_txt_text(file_path)
            else:
                # Should be unreachable given the extension check above.
                result.errors.append(f"No extractor implemented for '{extension}'.")
                return result
        except Exception as exc:  # noqa: BLE001 - we intentionally catch-all here
            # Any unexpected failure inside an extractor is converted into
            # a diagnostic error rather than crashing the whole pipeline.
            logger.exception("Unexpected error while parsing %s", file_path)
            result.errors.append(f"Unexpected parsing error: {exc}")
            return result

        result.warnings.extend(warnings)
        cleaned_text = self._normalize_whitespace(text)
        result.raw_text = cleaned_text
        result.char_count = len(cleaned_text)

        # --- Validate extraction quality ---------------------------------
        if result.char_count < MIN_VALID_TEXT_LENGTH:
            result.errors.append(
                "Extracted text is too short or empty. The file may be a "
                "scanned/image-based document with no selectable text, "
                "corrupted, or password-protected."
            )
            result.success = False
            return result

        result.success = True
        return result

    def parse_bytes(
        self, file_bytes: bytes, filename: str
    ) -> ParsedResume:
        """Parse resume content already loaded in memory (e.g. from an
        upload widget) by writing it to a temp file and delegating to
        `parse`. Kept separate from `parse` so callers with in-memory
        uploads (common in UI layers) don't need to manage temp files
        themselves.
        """
        import tempfile

        result = ParsedResume(file_path=filename)
        extension = Path(filename).suffix.lower()

        if extension not in SUPPORTED_EXTENSIONS:
            result.errors.append(
                f"Unsupported file type '{extension}'. "
                f"Supported types: {', '.join(SUPPORTED_EXTENSIONS.values())}"
            )
            return result

        if not file_bytes:
            result.errors.append("Provided file content is empty.")
            return result

        try:
            with tempfile.NamedTemporaryFile(
                suffix=extension, delete=False
            ) as tmp_file:
                tmp_file.write(file_bytes)
                tmp_path = tmp_file.name

            parsed = self.parse(tmp_path)
            # Preserve the original filename in the result for readability,
            # rather than exposing the temp file path to callers.
            parsed.file_path = filename
            return parsed
        except Exception as exc:  # noqa: BLE001
            logger.exception("Failed to parse in-memory file %s", filename)
            result.errors.append(f"Unexpected parsing error: {exc}")
            return result
        finally:
            try:
                if "tmp_path" in locals() and os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except OSError:
                pass  # Best-effort cleanup; not fatal.

    # -- Extraction backends ----------------------------------------------

    def _extract_pdf_text(self, file_path: str) -> tuple[str, list[str]]:
        """Extract text from a PDF, preferring pdfplumber (better layout
        handling) and falling back to PyPDF2 if pdfplumber fails or is
        unavailable. Returns (text, warnings).
        """
        warnings: list[str] = []

        if not _PDFPLUMBER_AVAILABLE and not _PYPDF2_AVAILABLE:
            raise RuntimeError(
                "No PDF parsing library available. Install pdfplumber or PyPDF2."
            )

        # --- Attempt 1: pdfplumber (primary) ---
        if _PDFPLUMBER_AVAILABLE:
            try:
                text_parts: list[str] = []
                with pdfplumber.open(file_path) as pdf:
                    if len(pdf.pages) == 0:
                        warnings.append("PDF has zero pages.")
                    for page_number, page in enumerate(pdf.pages, start=1):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                        except Exception as page_exc:  # noqa: BLE001
                            warnings.append(
                                f"pdfplumber failed on page {page_number}: {page_exc}"
                            )
                combined = "\n".join(text_parts).strip()
                if combined:
                    return combined, warnings
                warnings.append(
                    "pdfplumber extracted no text; attempting PyPDF2 fallback."
                )
            except Exception as exc:  # noqa: BLE001
                warnings.append(
                    f"pdfplumber failed to open/parse PDF ({exc}); "
                    "attempting PyPDF2 fallback."
                )

        # --- Attempt 2: PyPDF2 (fallback) ---
        if _PYPDF2_AVAILABLE:
            try:
                text_parts = []
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)

                    if getattr(reader, "is_encrypted", False):
                        try:
                            reader.decrypt("")  # try an empty password
                        except Exception:
                            raise RuntimeError(
                                "PDF is password-protected and could not be "
                                "decrypted."
                            )

                    for page_number, page in enumerate(reader.pages, start=1):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                        except Exception as page_exc:  # noqa: BLE001
                            warnings.append(
                                f"PyPDF2 failed on page {page_number}: {page_exc}"
                            )
                combined = "\n".join(text_parts).strip()
                return combined, warnings
            except Exception as exc:  # noqa: BLE001
                raise RuntimeError(f"PDF extraction failed with both backends: {exc}")

        # If pdfplumber produced nothing and PyPDF2 isn't available.
        return "", warnings

    def _extract_docx_text(self, file_path: str) -> tuple[str, list[str]]:
        """Extract text from a DOCX file, including paragraphs and table
        cell content (resumes often use tables for layout). Returns
        (text, warnings).
        """
        warnings: list[str] = []

        if not _DOCX_AVAILABLE:
            raise RuntimeError(
                "python-docx is not installed. Cannot parse DOCX files."
            )

        try:
            document = docx.Document(file_path)
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError(f"Could not open DOCX file (possibly corrupted): {exc}")

        text_parts: list[str] = []

        # Paragraphs (main body text).
        for paragraph in document.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Tables — many resumes place skills/experience in table layouts.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text_parts.append(cell.text.strip())

        if not text_parts:
            warnings.append("No paragraph or table text found in DOCX.")

        return "\n".join(text_parts), warnings

    def _extract_txt_text(self, file_path: str) -> tuple[str, list[str]]:
        """Extract text from a plain TXT file, trying a small set of
        common encodings before giving up. Returns (text, warnings).
        """
        warnings: list[str] = []
        encodings_to_try = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

        for encoding in encodings_to_try:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                if encoding != "utf-8":
                    warnings.append(
                        f"File was not UTF-8 encoded; read using '{encoding}'."
                    )
                return text, warnings
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as exc:  # noqa: BLE001
                raise RuntimeError(f"Could not read TXT file: {exc}")

        raise RuntimeError(
            "Could not decode TXT file with any supported encoding "
            f"({', '.join(encodings_to_try)})."
        )

    # -- Helpers ------------------------------------------------------------

    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        """Collapse excessive blank lines/spaces produced by extraction
        artifacts, while preserving paragraph structure for downstream
        section detection.
        """
        if not text:
            return ""

        # Normalize line endings.
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Collapse 3+ consecutive newlines into a max of 2 (one blank line).
        lines = [line.rstrip() for line in text.split("\n")]
        normalized_lines: list[str] = []
        blank_streak = 0

        for line in lines:
            if line.strip() == "":
                blank_streak += 1
                if blank_streak <= 1:
                    normalized_lines.append("")
            else:
                blank_streak = 0
                # Collapse internal multiple spaces/tabs to a single space.
                normalized_lines.append(" ".join(line.split()))

        return "\n".join(normalized_lines).strip()


# --- Convenience module-level function --------------------------------------

def parse_resume(file_path: str) -> ParsedResume:
    """Convenience wrapper for one-off parsing without instantiating the
    class explicitly.
    """
    return ResumeParser().parse(file_path)


if __name__ == "__main__":
    # Simple manual smoke test when run directly:
    #   python ml/resume_parser.py path/to/resume.pdf
    import sys

    if len(sys.argv) < 2:
        print("Usage: python resume_parser.py <path_to_resume_file>")
        sys.exit(1)

    parsed_result = parse_resume(sys.argv[1])
    print(f"Success: {parsed_result.success}")
    print(f"File type: {parsed_result.file_type}")
    print(f"Characters extracted: {parsed_result.char_count}")
    if parsed_result.warnings:
        print(f"Warnings: {parsed_result.warnings}")
    if parsed_result.errors:
        print(f"Errors: {parsed_result.errors}")
    print("\n--- Extracted Text Preview ---")
    print(parsed_result.raw_text[:500])